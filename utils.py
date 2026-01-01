import os
import json
from typing import List, Dict

import pypdf
import docx2txt
from langchain_openai import ChatOpenAI
from langchain_core.prompts import PromptTemplate
from langchain_core.output_parsers import StrOutputParser
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn




# =========================================================
# FILE PARSERS
# =========================================================

def parse_pdf(file) -> str:
    reader = pypdf.PdfReader(file)
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n".join(pages).strip()


def parse_docx(file) -> str:
    try:
        # Try python-docx first (more robust standard library)
        doc = Document(file)
        return "\n".join([p.text for p in doc.paragraphs]).strip()
    except Exception:
        try:
            # Fallback to docx2txt (handles some older XML formats better)
            file.seek(0)
            return docx2txt.process(file).strip()
        except Exception as e:
            return f"Error parsing DOCX: {str(e)}"


# =========================================================
# VERIFICATION AGENT (CRITIC)
# =========================================================

def verify_optimization(original_resume: str, optimized_resume: str, plan: List[Dict], api_key: str) -> List[Dict]:
    llm = get_llm(api_key)
    
    plan_text = "\n".join([
        f"- {item.get('type')}: {item.get('proposal') or item.get('proposed_text')}" 
        for item in plan
    ])

    prompt = """
    You are a Strict Code Auditor for resumes.
    
    GOAL: Verify if the requested changes were ACTUALLY applied to the resume.
    
    INPUT DATA:
    1. REQUESTED CHANGES (The Plan):
    {plan_text}
    
    2. OPTIMIZED RESUME (The Output):
    {optimized_resume}
    
    INSTRUCTIONS:
    - For EACH requested change in the plan, search the "OPTIMIZED RESUME" for evidence.
    - If the change is found (e.g. keyword added, bullet rewritten, OR NEW BULLET ADDED), mark as "VERIFIED".
    - NOTE: If the plan said "Rewrite bullet" but you find a NEW matching bullet, that IS A SUCCESS. Mark it "VERIFIED".
    - If the change is missing or hallucinated, mark as "FAILED".
    - Be strict about content presence, but flexible about placement.

    OUTPUT JSON format:
    [
        {{
            "proposal": "Add 'Agile' to skills",
            "status": "VERIFIED" | "FAILED",
            "evidence": "Found 'Agile Methodologies' in Skills section",
            "notes": "Exact match found"
        }}
    ]
    """
    
    chain = PromptTemplate(
        input_variables=["plan_text", "optimized_resume"],
        template=prompt
    ) | llm | StrOutputParser()
    
    print("DEBUG: Running Verification Agent...")
    try:
        raw = chain.invoke({
            "plan_text": plan_text, 
            "optimized_resume": optimized_resume
        })
        # Clean markdown
        if "```" in raw:
            raw = raw.replace("```json", "").replace("```", "").strip()
        return json.loads(raw)
    except Exception as e:
        print(f"VERIFICATION ERROR: {e}")
        return []

# =========================================================
# LLM
# =========================================================

def get_llm(api_key: str) -> ChatOpenAI:
    return ChatOpenAI(
        model="gpt-4o",
        openai_api_key=api_key,
        temperature=0.0,
        model_kwargs={"seed": 42}
    )


# =========================================================
# ATS ANALYZER (ENTERPRISE-REALISTIC)
# =========================================================

def analyze_resume(resume_text: str, jd_text: str, api_key: str) -> Dict:
    llm = get_llm(api_key)

    prompt = """
You are simulating a REAL enterprise ATS ranking system
used by companies running Workday, Greenhouse, or iCIMS.

========================
REAL ATS RULES (STRICT)
========================
• ATS systems RANK candidates; they do not reason, infer, or compensate
• Only REQUIRED qualifications affect eligibility
• Preferred qualifications influence ranking but NEVER block eligibility
• Role-intent mismatch applies ONLY when the ROLE TYPE or TITLE mismatches
  (e.g., Data Engineer vs Frontend Engineer)
• Senior overqualification causes a normalization penalty
• Noise (irrelevant skills) slightly dilutes ranking but never disqualifies

========================
SECTOR EXPERIENCE RULE (CRITICAL)
========================
• Sector experience (e.g., Agriculture, Healthcare, Energy) MUST be treated as:
  - A BASIC QUALIFICATION ONLY IF the JD explicitly states it as REQUIRED
    Examples:
      - "Agriculture experience required"
      - "Must have healthcare domain experience"
      - "Energy or chemical industry background required"

• If sector experience IS REQUIRED:
  - Search the ENTIRE resume for explicit sector keywords
  - If sector keywords are MISSING:
      → Candidate FAILS eligibility (hard block)
      → meets_minimum_requirements = false
      → Add blocking issue explaining missing required sector
  - If sector keywords are PRESENT:
      → No penalty and no bonus (binary gate only)

• If sector experience is NOT explicitly required:
  - Sector has ZERO impact on eligibility
  - Sector has ZERO impact on ranking
  - DO NOT penalize, downgrade, or infer intent mismatch

========================
STRICT KEYWORD MATCHING
========================
• STRICT LITERAL MATCHING with GLOBAL SEARCH:
   - Verify the presence of EXACT keywords from the JD in the Resume
   - SEARCH ENTIRE TEXT (Experience, Skills, Projects, Summary)
   - AGGREGATE EVIDENCE across sections
   - Do NOT infer synonyms
   - Do NOT infer "Agile" from "Scrum"

• MATCHING ANALYSIS (MANDATORY):
    - For EVERY requirement line in the JD:
    FORBIDDEN FALSE POSITIVES (TOP PRIORITY):
    - Requirement: "Automated Testing" (Unit, Integration, E2E)
      - INVALID EVIDENCE: "CI/CD", "Pipelines", "DevOps", "Jenkins".
      - VALID EVIDENCE: "JUnit", "Selenium", "PyTest", "Jest", "Cypress", "Mocking", "TDD".
      - RULE: If you only find "CI/CD" -> Status: MISSING. (Do not mark Full. Do not mark Partial).
      - REASONING: Building a pipeline is not the same as writing a test.

    - Requirement: "Agile and DevOps"
      - RULE: You MUST find explicit keywords for BOTH.
      - FOR "Agile": You MUST find "Agile", "Scrum", "Kanban", or "SAFe".
        - INVALID: "TDD", "BDD", "Jira", "Sprints", "Standups", "Iterative". (These are practices, not the methodology itself).
      - IF "Agile" keyword is missing -> Status: PARTIAL (Note: "Missing explicit keyword 'Agile'").
      - IF "DevOps" is missing -> Status: PARTIAL.
      - "Full" status requires explicit keywords for BOTH.

    ANALYSIS INSTRUCTIONS (Follow in Order):
    
    0. EXPERIENCE CALCULATION (MANDATORY):
       - Extract Candidate's TOTAL Years of Experience from work history.
       - Extract JD's REQUIRED Years of Experience.
       - Calculate "Seniority Status":
         - "Underqualified" (Candidate < Min JD)
         - "Match" (Candidate within range)
         - "Overqualified" (Candidate > Max JD + 2 years)
       - RETURN these values in the JSON output.

    1. ROLE ALIGNMENT CHECK (CRITICAL - DO THIS FIRST):
       - Compare the Candidate's CURRENT TITLE and PRIMARY SKILL SET vs the JOB TITLE.
       - SW vs DATA: If Job is "Software Engineer" (Application Dev, Java, APIs) and Resume is "Data Engineer" (ETL, Spark, Pipelines) -> STATUS: "Mismatch".
       - This is a PRIMARY FILTER. Be strict.

    2. SOFT SKILLS (STRICT DEFAULT):
       - For abstract requirements ("Communication", "Teamwork", "Time Management", "Leadership"):
       - DEFAULT STATUS: "Partial" (or Missing).
       - YOU MAY ONLY GRANT "FULL" IF:
         - The resume contains VERBATIM evidence or specific examples (e.g. "Led team of 5", "Presented to stakeholders").
         - Vague claims ("Good communicator") are insufficient -> Status: Partial.
       - DO NOT infer soft skills from "having a job".
    
    2. SIMPLIFIED MATCHING LOGIC (For Requirements):
    
    1. "AND" LOGIC (CRITICAL):
       - Requirement: "Experience with A AND B" (e.g. "Agile and DevOps", "Java and MVC")
       - Status: 
         - FULL: Both A and B are explicitly found.
         - PARTIAL: Only one is found.
         - MISSING: Neither is found.
       - Note: "Found A. Missing B."
       - EXAMPLE: "Agile and DevOps"
         - Must find "Agile" (or Scrum/Kanban) AND "DevOps" (or CI/CD/Pipelines).
         - Finding "DevSecOps" proves DevOps, but NOT Agile. -> Status: Partial.

    2. "LIST/INCLUDING" LOGIC (OVERRIDES "OR"):
       - Trigger: Requirement contains "including", "such as", "e.g.", or a comma-separated list of tools.
       - RULE: "Full" status requires ALL listed items to be present, even if "or" is used.
       - REASONING: "Including A, B, or C" implies breadth of experience.
       - IF N < 5 items:
         - FULL: All items found.
         - PARTIAL: Any item missing. (e.g. Found React, but missing Node.js -> Partial).
         - MISSING: All items missing.
       - IF N >= 5 items:
         - Apply >80% threshold for Full.

    3. "SIMPLE OR" LOGIC (Only for short, non-list choices):
       - Requirement: "Experience with Java or Python" (No "including", no long list).
       - Status: FULL if ANY item is found.

    4. EVIDENCE QUALITY & QUOTE SELECTION (CRITICAL):
       - RULE: You are FORBIDDEN from quoting a line that mentions "years" unless it ALSO mentions the specific language (e.g. "Java").
       - EXCEPTION FOR SOFT SKILLS: 
         - Requirements like "learning new technologies" require evidence of *change*, *migration*, or *adoption*.
         - VALID: "Migrated legacy app to Cloud", "Adopted Go for new services", "Researched and implemented X".
         - INVALID: "Iterated on features", "Worked in Agile team", "Delivered projects".
       - FINAL CHECK (Technical Only): Does your selected quote contain the specific Tech Keyword?

     6. STRICT KEYWORD MATCHING (ZERO INFERENCE):
        - You are an ATS machine, NOT a human recruiter.
        - You CANNOT infer skills. 
        - If the JD asks for "MVC" and the resume says "Spring Boot" but NOT "MVC":
          - Status: PARTIAL (or Missing).
          - Note: "Found 'Spring Boot'. Missing explicit keyword 'MVC'."
        - REASON: The goal is to force the user to add the EXACT keyword to their resume.
        - "Concept met" is NOT enough. The *Keyword* must be present.

    7. SOFT SKILL INTEGRITY (ANTI-PATTERNS):
        - REQUIREMENT: "Ability to learn new technologies" or "Adaptability".
        - INVALID EVIDENCE (Status: MISSING):
          - "Implemented X using Y" (This is just doing the job).
          - "Worked in Agile environment".
          - "Iterated on features".
        - VALID EVIDENCE (Status: FULL):
          - "Self-taught Python..."
          - "Migrated legacy monolithic app to microservices..." (Shows adaptation).
          - "Researched and adopted..."
          - "PoC" / "Proof of Concept".
        - IF NO EXPLICIT EVIDENCE OF LEARNING/MIGRATION IS FOUND -> MARK AS MISSING.

    8. STATUS INTEGRITY CHECK (MANDATORY):
       - Look at your own "Notes" field.
       - IF "Notes" contains the word "Missing" followed by a specific tool/skill name:
         - YOU MUST SET STATUS TO "PARTIAL" (or Missing).
         - YOU ARE FORBIDDEN FROM SETTING STATUS TO "FULL".
       - Example: Notes="Found React. Missing Node." -> Status MUST be "Partial".
       - Example: Notes="Missing explicit mention of Agile" -> Status MUST be "Partial".

        STATUS INTEGRITY RULE (ABSOLUTE):
       - Look at your own "Notes" field.
       - IF "Notes" contains the word "Missing" or listing missing items:
         - The "Status" MUST be "Partial".
         - You are FORBIDDEN from marking it "Full".
       - Example: Notes="Found React. Missing Node.js." -> Status: "Partial" (NOT Full).
       - Example: Notes="Missing explicit mention of Agile" -> Status: "Partial".

    7. EVIDENCE INTEGRITY RULE (GLOBAL):
       - You CANNOT use purely generic terms to satisfy specific requirements.
       - "CI/CD pipelines" alone is NOT evidence for "Automated Testing".
       - VALID EVIDENCE: Unit, Integration, E2E, Selenium, Jest, PyTest, JUnit, TestNG, Cypress.
       - ALSO VALID: Specific testing tools like SonarQube (Static Analysis), Postman (API Testing).
       - "Cloud technologies" is NOT evidence for "AWS".
       - IF the specific keyword is not in the text, mark as MISSING.
       

    MANDATORY NOTES FORMAT (Show your math):
        "Matched: 4/7 (57%). Found: 'Lambda', 'API Gateway', 'DynamoDB', 'S3'. Missing: 'SNS', 'SQS', 'EventBridge'."
        
        FOR "OR" REQUIREMENTS:
        - Even if Status is "Full", you MUST list the missing options.
        - Example: "Found: 'React'. Missing: 'Node.js', 'AngularJS'."

    9. HYBRID REQUIREMENT EXTRACTION (CRITICAL):
       - If a single line in the JD says: "Experience with X required; Y and Z are a plus."
       - SPLIT THIS into two separate entries:
         1. Basic Qual: "Experience with X"
         2. Preferred Qual: "Experience with Y and Z"
       - DO NOT LUMP THEM TOGETHER. This ensures we track the "plus" items separately.

    10. CAPTURE ALL "INTANGIBLE" PREFERENCES:
        - You MUST extract lines starting with "Interest in...", "Passion for...", "Willingness to...", "Ability to...".
        - Do not skip them. They are valid Preferred Qualifications.
        - Example: "General interest in Sports" -> Extract as Preferred Qual.

    11. COPY-PASTE REQUIREMENT TEXT (MANDATORY):
        - You are FORBIDDEN from creating your own requirement names.
        - You MUST COPY AND PASTE the exact text from the JD.
        - BAD: "AWS List"
        - GOOD: "Experience with Cloud AWS-hosted applications"
        - BAD: "Agile and DevOps"
        - GOOD: "Experience working in an agile environment"
        - RULE: If you summarize, the user will fail the interview. Do not help them by summarizing.

    12. STANDALONE PREFERRED QUALIFICATIONS:
        - If a line says "...is a plus" or "...are a plus" or "Preferred:", EXTRACT IT as a Preferred Qual.
        - Do not ignore it.
        - Example: "Experience with Kotlin Multiplatform is a plus" -> Extract to Preferred.




========================
PRIMARY QUESTION
========================
How intentionally aligned is this resume to THIS role
based ONLY on JD-stated requirements?

========================
JOB DESCRIPTION
========================
{jd_text}

========================
RESUME
========================
{resume_text}

========================
OUTPUT FORMAT (VALID JSON ONLY)
========================
{{
  "eligibility_check": {{
    "meets_minimum_requirements": true,
    "blocking_issues": []
  }},
  "role_intent_alignment": "Low | Medium | High",
  "required_skill_signals": {{
    "primary": [],
    "secondary": []
  }},
  "preferred_skill_signals": [],
  "cloud_alignment": {{
    "primary_cloud_match": "AWS | Azure | Mixed",
    "serverless_experience": true
  }},
  "delivery_signals": {{
    "automated_testing": true,
    "ci_cd": true
  }},
  "noise_or_dilution_signals": [],
  "keyword_gap_analysis": {{
    "missing_jd_keywords": [],
    "concept_mismatches": []
  }},
  "basic_qualification_match": [
    {{
      "requirement": "VERBATIM text from JD (Do not summarize)",
      "match_status": "Full | Partial | Missing",
      "evidence_quote": "Exact substring from resume (or empty)",
      "notes": "List exactly what is Found vs Missing"
    }}
  ],
  "role_alignment": {{
    "match_status": "Aligned" | "Mismatch" | "Partial",
    "reason": "Explanation of misalignment (e.g. 'Resume is Data Engineer, JD is Software Engineer')"
  }},
  "preferred_qualification_match": [
    {{
      "requirement": "VERBATIM text from JD (Do not summarize)",
      "match_status": "Full | Partial | Missing",
      "evidence_quote": "Exact substring from resume (or empty)",
      "notes": "List exactly what is Found vs Missing"
    }}
  ],
  "seniority_normalization_penalty": 0,
  "role_mismatch_penalty": 0,
  "role_intent_penalty": 0,
  "overall_relevance_score": 0,
  "ats_ranking_band": "Low | Medium | High | Very High",
  "recruiter_summary": ""
}}

========================
SCORING BOUNDS (MATH ONLY)
========================
1. Base Score: Start at 100
2. Penalties:
   - For each MISSING Basic Req: -15
   - For each PARTIAL Basic Req: -8
   - For ROLE MISMATCH ("Mismatch"): -20 (Critical Penalty)
   - For SENIORITY MISMATCH (Overqualified): -10
   - For INTENT MISMATCH: -5
3. Final Score = 100 - (Sum of Penalties)
   - Minimum Score: 0
   - CRITICAL: You must CALCULATE the final number (e.g. 85). Do NOT return a formula string like "100 - 15". Return the INTEGER value.
"""

    chain = PromptTemplate(
        input_variables=["resume_text", "jd_text"],
        template=prompt
    ) | llm | StrOutputParser()

    raw = chain.invoke({
        "resume_text": resume_text,
        "jd_text": jd_text
    }).strip()

    if raw.startswith("```"):
        raw = raw.replace("```json", "").replace("```", "").strip()

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        raise RuntimeError(f"ATS analyzer returned invalid JSON:\n{raw}")


def verify_analysis_integrity(analysis: Dict, resume_text: str, jd_text: str, api_key: str) -> Dict:
    """
    Second-pass Auditor to verify the initial ATS analysis.
    Checks for:
    1. False Negatives: Did we mark something as 'Missing' that is actually there?
    2. False Positives: Did we accept generic evidence for a specific tool?
    """
    llm = get_llm(api_key)
    
    # Extract only the critical parts to verify to save tokens
    to_verify = {
        "basic_qualification_match": analysis.get("basic_qualification_match", []),
        "preferred_qualification_match": analysis.get("preferred_qualification_match", []),
        "role_alignment": analysis.get("role_alignment", {})
    }
    
    prompt = """
    You are a Senior QA Auditor for an ATS algorithm.
    
    Your job is to REVIEW the preliminary analysis of a resume and CORRECT any mistakes.
    The first analyzer (ATS Bot) is strict and sometimes misses things. You must fix this.
    
    INPUTS:
    1. Resume Text (Source of Truth):
    {resume_text}

    2. Preliminary Analysis (Draft to be graded):
    {analysis_json}
    
    YOUR TASKS (Apply to ALL requirements):
    
    1. AUDIT "MISSING" ITEMS (Search Deeply):
       - If a skill is marked "Missing" (e.g. "SNS", "SQS"), SCAN THE RESUME.
       - Is the exact word there? 
       - Is a clear synonym there? (e.g. "Postgres" for "SQL")
       - IF FOUND: YOU MUST UPGRADE status to "Full" or "Partial" and update the Note.
       - Note format: "Auditor Correction: Found 'X' in [Section]."
       
    2. AUDIT "FULL" ITEMS (Verify Evidence):
       - specific keywords must be present.
       - COMPOUND CHECK: If requirement is "Agile AND DevOps", evidence MUST show BOTH.
       - LIST CHECK (CRITICAL): If requirement lists "X, Y, Z" (e.g. AWS List), ALL MUST BE PRESENT for "Full".
       - "Agile" Specifics:
         - You CANNOT infer Agile from "Scrum", "Sprints", "Jira", "TDD", "BDD", or "CI/CD" alone.
         - If the JD says "Agile environment", you must find the word "Agile" explicitly listed.
         - IF NO EXPLICIT "Agile" keyword is found, Downgrade to "Partial" or "Missing".
         - Note: "Inferred Agile from context/TDD" -> Status: Partial.
       - If evidence is weak or hallucinated, Downgrade to "Partial" or "Missing".
       
    3. SOFT SKILL AUDIT (Critical):
       - Requirement: "Ability to learn new technologies" or similar.
       - Evidence: "Implemented BDD/TDD..." -> INVALID. (Mark as Missing).
       - Evidence: "Migrated..." or "Self-taught..." -> VALID.
       - Verify if the Role Alignment assessment is fair match.

    4. HYBRID REQUIREMENTS (Basic + Bonus):
       - If a unique item is listed as "a plus", "preferred", "ideally", or "nice to have" WITHIN a Basic Qualification:
       - AND that item is missing from the resume:
       - YOU MUST listed it in "Missing: ..."
       - Example: Req="Java required; Kafka is a plus". Found Java (Full). Note: "Found Java. Missing: Kafka." 

    5. VERBATIM REQUIREMENT NAMES (MANDATORY):
       - You MUST preserve the exact requirement text from the input.
       - DO NOT summarize (e.g. changing "Experience with AWS..." to "AWS List").
       - If you change the Requirement Name, the system breaks. KEEP IT EXACT.

    MANDATORY NOTE FORMATTING:
    - If Status is 'Partial' or 'Missing', or if it is 'Full' but one of many options:
    - YOU MUST END THE NOTE WITH: "Missing: item1, item2, item3".
    - This is REQUIRED for the system to generate the gap analysis.
    - Example: "Auditor Correction: Found React. Missing: Node.js, Typescript."

    RETURN FORMAT (STRICT JSON ONLY):
    Result must be valid JSON strictly following this structure.
    YOU MUST RETURN THE *COMPLETE* LIST OF REQUIREMENTS (All items from input, even if unchanged).
    Do NOT return a partial match list.
    Structure:
    {{
      "basic_qualification_match": [ ...all items... ],
      "preferred_qualification_match": [ ...all items... ],
      "role_alignment": {{...}}
    }}
    DO NOT ADD COMMENTARY OUTSIDE THE JSON.
    """
    
    chain = PromptTemplate(
        input_variables=["resume_text", "analysis_json"],
        template=prompt
    ) | llm | StrOutputParser()
    
    print("DEBUG: Running Analysis Auditor...")
    try:
        raw = chain.invoke({
            "resume_text": resume_text,
            "analysis_json": json.dumps(to_verify, indent=2)
        })
        print(f"DEBUG RAW AUDITOR RESPONSE: {raw[:500]}...") # Print first 500 chars
        
        if "```" in raw:
            raw = raw.replace("```json", "").replace("```", "").strip()
            
        corrections = json.loads(raw)
        
        # Merge corrections back into original analysis
        final_analysis = analysis.copy()
        final_analysis.update(corrections)
        
        # --- PYTHON HARD CHECK: AGILE ENFORCEMENT ---
        # The LLM is too lenient. We MUST enforce keyword presence for "Agile".
        # 1. Use Regex with Word Boundaries to avoid "xp" matching "experience" or "safe" matching "thread-safe".
        import re
        resume_lower = resume_text.lower()
        # Strictly accepted Agile keywords
        agile_patterns = [
            r'\bagile\b', 
            r'\bscrum\b', 
            r'\bkanban\b', 
            r'\bsafe\b', # Scaled Agile Framework - might match 'safe', but strictly as a whole word is rarer in non-tech context
            r'\bxp\b',   # Extreme Programming - boundary checks prevent 'experience' match
            r'\bextreme programming\b'
        ]
        
        # Remove "safe" if it risks false positives with "safe environment"? 
        # Actually "thread-safe" has a hyphen, which is a boundary. So \bsafe\b matches "safe" in "thread-safe".
        # Let's be STRICT: Only Agile, Scrum, Kanban, Extreme Programming.
        # "SAFe" is usually capitalized, but we are lowering. 
        # To be safe (pun intended), we remove "safe" and "xp" to avoid ANY ambiguity.
        # If they practice SAFe, they usually verify "Agile" too.
        
        strict_patterns = [r'\bagile\b', r'\bscrum\b', r'\bkanban\b', r'\bextreme programming\b', r'\bscrumban\b']
        
        has_agile_keyword = any(re.search(p, resume_lower) for p in strict_patterns)
        
        if not has_agile_keyword:
            # Check Basic Quals for "Agile" requirement
            for item in final_analysis.get("basic_qualification_match", []):
                req_lower = item.get("requirement", "").lower()
                if "agile" in req_lower:
                    # If LLM marked it Full (case-insensitive), Downgrade it.
                    current_status = str(item.get("match_status", "")).lower()
                    if current_status == "full":
                        item["match_status"] = "Partial"
                        item["notes"] = "Auto-Auditor: No explicit 'Agile' keyword found. Inferred methodology is NOT accepted."
                        print(f"DEBUG: Auto-Downgraded Agile requirement: {item['requirement']}")

            # Check Preferred Quals
            for item in final_analysis.get("preferred_qualification_match", []):
                req_lower = item.get("requirement", "").lower()
                if "agile" in req_lower:
                    current_status = str(item.get("match_status", "")).lower()
                    if current_status == "full":
                        item["match_status"] = "Partial"
                        item["notes"] = "Auto-Auditor: No explicit 'Agile' keyword found."
        
        # --- PYTHON POST-PROCESSING: AGGREGATE GAPS ---
        # The Auditor fixes the notes, but we need to update the top-level summary list
        # based on those validated notes.
        missing_keywords = set()
        
        # Scan basic quals
        for item in final_analysis.get("basic_qualification_match", []):
            notes = item.get("notes", "")
            if "Missing" in notes:
                # Extract words after 'Missing'
                # Simple heuristic: Look for 'Missing: 'X', 'Y''
                # This doesn't need to be perfect, just better than empty.
                parts = notes.split("Missing", 1)
                if len(parts) > 1:
                    missing_text = parts[1].strip(": .-")
                    # Clean up: "SNS", "SQS" -> ["SNS", "SQS"]
                    gaps = [g.strip(" '\",.") for g in missing_text.split(",")]
                    missing_keywords.update([g for g in gaps if g])

        # Scan preferred quals
        for item in final_analysis.get("preferred_qualification_match", []):
             notes = item.get("notes", "")
             if "Missing" in notes:
                parts = notes.split("Missing", 1)
                if len(parts) > 1:
                    missing_text = parts[1].strip(": .-")
                    gaps = [g.strip(" '\",.") for g in missing_text.split(",")]
                    missing_keywords.update([g for g in gaps if g])
        
        # Update the summary dictionary
        if "keyword_gap_analysis" not in final_analysis:
            final_analysis["keyword_gap_analysis"] = {}
            
        final_analysis["keyword_gap_analysis"]["missing_jd_keywords"] = list(missing_keywords)
        
        # --- PYTHON POST-PROCESSING: ROLE MISMATCH AS RISK ---
        # If the role is a mismatch, this is the BIGGEST risk. It should be in the Risk Signals box.
        role_status = final_analysis.get("role_alignment", {}).get("match_status", "Aligned")
        role_intent = final_analysis.get("role_intent_alignment", "High")
        
        if role_status == "Mismatch" or role_intent == "Low":
            final_analysis["noise_or_dilution_signals"] = risks
        
        # --- PYTHON POST-PROCESSING: RECALCULATE SCORE ---
        # The LLM's score is often hallucinated. We must calculate it based on the VERIFIED evidence.
        return calculate_verified_score(final_analysis)
        
    except Exception as e:
        print(f"AUDIT ERROR: {e}")
        # FAIL SAFE: Even if Auditor crashes, we MUST apply the strict scoring logic
        # to the original analysis to prevent 100/100 scores.
        print("DEBUG: Applying fallback scoring to original analysis.")
        return calculate_verified_score(analysis)

def calculate_verified_score(analysis: Dict) -> Dict:
    """
    Enterprise ATS Scoring Engine (Strict).
    Simulates Workday/Greenhouse ranking logic.
    """
    # RULE 2: Score Ceiling (Absolute)
    # 100 is impossible. Hard Absolute Cap = 94.
    score = 94
    meets_min_reqs = True
    perfect_preferred = True # Track if we hit the "88 Cap" condition
    
    # Soft Skill Keywords (for Rule 4)
    soft_skill_keywords = ["communication", "teamwork", "leadership", "time management", "adaptability", "collaborat", "interpersonal"]

    # 1. BASIC QUALIFICATIONS (The Gatekeepers)
    # Rule 3: Missing -> -20, Partial -> -10
    for item in analysis.get("basic_qualification_match", []):
        status = str(item.get("match_status", "Full")).lower()
        req_text = str(item.get("requirement", "")).lower()
        
        # RULE 4: Soft Skill Inflation Fix
        # If it's a soft skill and marked Full, be skeptical.
        is_soft_skill = any(k in req_text for k in soft_skill_keywords)
        if is_soft_skill and status == "full":
            # Downgrade to Partial automatically unless we trust the LLM fully (we don't)
            # Apply -3 penalty as per Rule 4
            score -= 3
            item["match_status"] = "Partial (Auto-Downgraded Soft Skill)"
            status = "partial" 
        
        if status == "missing":
            meets_min_reqs = False
            score -= 20
        elif status == "partial":
            score -= 10
            
    # 2. PREFERRED QUALIFICATIONS (The Differentiators)
    # Rule 5: Missing -> -5, Partial -> -3
    for item in analysis.get("preferred_qualification_match", []):
        status = str(item.get("match_status", "Full")).lower()
        
        if status == "missing":
            score -= 5
            perfect_preferred = False
        elif status == "partial":
            score -= 3
            perfect_preferred = False
            
    # 3. ROLE ALIGNMENT (Rule 7)
    role_data = analysis.get("role_alignment", {})
    role_status = str(role_data.get("match_status", "Aligned")).lower()
    
    if "mismatch" in role_status:
        score -= 15
    elif "partial" in role_status:
        score -= 8

    # 4. SENIORITY NORMALIZATION (Rule 8)
    seniority = str(analysis.get("seniority_status", "Match")).lower()
    
    if "overqualified" in seniority:
        score -= 10 
    elif "underqualified" in seniority:
        score -= 15
        
    # 5. INDUSTRY EXPERIENCE (Rule 6)
    # Check noise signals for industry mismatch
    risks = analysis.get("noise_or_dilution_signals", [])
    for r in risks:
        if "industry" in str(r).lower():
            score -= 5 # Soft dilution

    # RULE 2: Conditional Cap
    # If ANY preferred qualification is Partial or Missing -> cap score at 88.
    if not perfect_preferred:
        score = min(score, 88)

    # FINAL BOUNDS
    score = max(0, min(94, score))
    
    # Update the dictionary in place
    analysis["overall_relevance_score"] = score
    analysis["meets_minimum_requirements"] = meets_min_reqs
    
    return analysis
    
    # Update Ranking Band
    if final_score >= 90:
        analysis["ats_ranking_band"] = "Very High"
    elif final_score >= 75:
            analysis["ats_ranking_band"] = "High"
    elif final_score >= 60:
            analysis["ats_ranking_band"] = "Medium"
    else:
            analysis["ats_ranking_band"] = "Low"

    return analysis


# =========================================================
# ATS RESUME REWRITER (INTENT-LOCKED)
# =========================================================

def rewrite_resume(
    resume_text: str,
    jd_text: str,
    optimization_plan: List[Dict],  # NEW: List of approved changes
    role_intent: str,
    seniority_penalty: int,
    api_key: str
) -> str:

    llm = get_llm(api_key)

    # Separate plan into Edits vs Additions for clarity
    edits = []
    additions = []
    
    # CONFLICT RESOLUTION: Merge overlapping edits for the same section
    # e.g. User wants to Pivot to GCP (Rewrite) AND Add Bedrock (Rewrite/Add)
    # We must merge these instructions or one will overwrite the other.
    
    from collections import defaultdict
    section_map = defaultdict(list)
    for i, item in enumerate(optimization_plan):
        section_map[item.get('target_section', 'General')].append(i)
        
    indices_to_remove = []
    
    for section, indices in section_map.items():
        if len(indices) > 1 and section != "General":
            # Check if we have multiple rewrites
            rewrites = [optimization_plan[i] for i in indices if optimization_plan[i].get('type') == 'FULL_ROLE_REWRITE']
            if len(rewrites) > 1:
                # CONFLICT DETECTED
                print(f"DEBUG: resolving conflict for {section}")
                
                # Identify the "Pivot" (Platform Change) vs "Additions"
                # Heuristic: Pivot usually has 'optimization_strategy'='PIVOT' or covers cloud terms
                
                # We will ask the LLM to merge them
                content_a = rewrites[0].get('proposal', [])
                content_b = rewrites[1].get('proposal', [])
                
                reason_a = rewrites[0].get('reason', '')
                reason_b = rewrites[1].get('reason', '')
                
                merge_prompt = f"""
                You are a Code/Text Merger.
                I have two versions of a Resume Role Description for '{section}'.
                
                VERSION A (Focus: {reason_a}):
                {content_a}
                
                VERSION B (Focus: {reason_b}):
                {content_b}
                
                TASK: Merge them into ONE coherent list.
                RULES:
                1. If one version is a 'Platform Pivot' (e.g. AWS->GCP), USE THAT stack as the base.
                2. If the other version adds a specific tool (e.g. Bedrock), INSERT that bullet into the Base.
                3. Remove Conflicts: Do NOT mention the Old Platform (e.g. AWS) if we pivoted to New (GCP), unless it's a multi-cloud role.
                4. Output ONLY the merged list of bullet points as a JSON List of Strings.
                """
                
                try:
                    from langchain.schema import HumanMessage
                    res = llm.invoke([HumanMessage(content=merge_prompt)])
                    content = res.content
                    if "```json" in content:
                        content = content.split("```json")[1].split("```")[0]
                    merged_list = json.loads(content)
                    
                    # Update First Item
                    optimization_plan[indices[0]]['proposal'] = merged_list
                    optimization_plan[indices[0]]['reason'] = f"Merged: {reason_a} + {reason_b}"
                    
                    # Mark others for deletion
                    for idx in indices[1:]:
                        indices_to_remove.append(idx)
                        
                except Exception as e:
                    print(f"Merge failed: {e}")
                    # Fallback: Keep both sequential (risky but better than crash)

        else:
            # Check for multiple APPEND_BULLETS
            appends = [optimization_plan[i] for i in indices if optimization_plan[i].get('type') == 'APPEND_BULLETS']
            if len(appends) > 1:
                print(f"DEBUG: Consolidating {len(appends)} APPEND items for {section}")
                
                # Consolidate bullets
                merged_bullets = []
                reasons = []
                
                for item in appends:
                    content = item.get('proposal', [])
                    if isinstance(content, list):
                        merged_bullets.extend(content)
                    elif isinstance(content, str):
                        merged_bullets.append(content)
                    
                    reasons.append(item.get('reason', ''))
                
                # Update First Item
                optimization_plan[indices[0]]['proposal'] = merged_bullets
                optimization_plan[indices[0]]['reason'] = " + ".join(reasons)
                
                # Mark others for deletion
                for idx in indices[1:]:
                    if optimization_plan[idx].get('type') == 'APPEND_BULLETS':
                        indices_to_remove.append(idx)

    # Remove merged duplicates
    optimization_plan = [item for i, item in enumerate(optimization_plan) if i not in indices_to_remove]

    for item in optimization_plan:
        type_ = item.get('type', 'CHANGE').upper()
        content = item.get('proposal') or item.get('proposed_text', 'No details')
        reason = item.get('reason', 'N/A')
        target = item.get('target_section', 'General')
        
        # New robust formatting
        if isinstance(content, list):
            formatted_content = "\n" + "\n".join([f"    • {c}" for c in content])
        else:
            formatted_content = str(content)
            
        line = f"- [TARGET: {target}] {type_}: {formatted_content} (Reason: {reason})"
        
        if type_ in ["FORCE_INSERTION", "ADD_SKILL"]:
            additions.append(line)
        else:
            edits.append(line)

    plan_text = ""
    if edits:
        plan_text += "*** EDITS TO EXISTING CONTENT ***\n" + "\n".join(edits) + "\n\n"
    if additions:
        plan_text += "*** MANDATORY ADDITIONS (MUST INSERT) ***\n" + "\n".join(additions)

    prompt = """
You are a Resume Optimizer tuned to ENTERPRISE ATS ranking behavior.

THIS IS A STRICT, CONTROLLED REWRITE.

NON-NEGOTIABLE RULES:
1. Preserve section order, company names, and dates EXACTLY.
2. TITLE EXCEPTION: You MAY rewrite the Professional Headline/Title ONLY if the Plan explicitly requests it (e.g. to fix Role Mismatch).
3. TARGET LOCK: Only modify the section specified in [TARGET: ...]. Do not touch other sections.

    You are a Resume Editor. Your ONLY job is to execute the following list of approved edits.

    *** APPROVED OPTIMIZATION PLAN (HIGHEST PRIORITY) ***
    {plan_text}
    *****************************************************

    EXECUTION RULES:
    1. FORCE INSERTION: You MUST apply every single edit listed above.
    2. OVERRIDE INTEGRITY: If the plan says "Rewrite bullet to include X", you MUST do it.
    3. FALLBACK: If you cannot find a specific existing bullet to rewrite, YOU MUST ADD A NEW BULLET with the content.
    4. PRESERVE STRUCTURE: Keep the rest of the resume exactly as is. Only modify the targeted sections.
    5. FULL CONTEXT REWRITE: If type is 'FULL_ROLE_REWRITE', replace ALL bullets under that specific Role/Company with the new list provided. Do NOT keep the old bullets.
    6. APPEND_BULLETS: If type is 'APPEND_BULLETS', find the specific Role/Company and ADD these new bullets to the END of the bullet list. Maintain all existing bullets.
    
    GAP FILLING (SECONDARY):
    - If the plan is empty or vague, ensure these terms are present:
    
    1. Formatting Rules:
2. Each role format: Company Name (Line 1) -> Date (Line 2) -> Title (Line 3) -> Bullets (Lines 4+)
3. STRICT BULLET RULE: The experience description bullets MUST start with "•". 
4. EXEMPTION: Do NOT put bullets on the Company Name, Date line, or Job Title line.
5. Do NOT invent experience.
6. Do NOT remove senior credibility
7. DOMAIN INTEGRITY: 
   - Do NOT change the actual industry.
   - DO optimize phrasing: Instead of "patient data", use "high-volume sensitive data".
   
8. HEADER PROTECTION (CRITICAL):
   - NEVER add text to the same line as a Company Name or Date.
   - New bullets must be on their own lines starting with "•".
   
9. SKILLS FORMATTING:
   - For Technical Skills, use the format: **Category Name:** Skill 1, Skill 2...
   - Ensure each category is on a NEW bullet line.
   - Do NOT append loose words at the bottom.
   
INTENT RULES:
• Resume must read as INTENTIONALLY aligned to THIS role
• Primary skills appear in summary and first bullets
• Secondary skills appear later
• Noise skills are pushed down or softened

JOB DESCRIPTION:
{jd_text}

TARGET ROLE INTENT:
{role_intent}

AUTHORITATIVE RESUME:
{resume_text}

OUTPUT:

OUTPUT:
Full resume text only
No explanations
No markdown
"""

    chain = PromptTemplate(
        input_variables=[
            "resume_text",
            "jd_text",
            "plan_text",
            "role_intent"
        ],
        template=prompt
    ) | llm | StrOutputParser()

    print(f"DEBUG: Rewriter Plan Text:\n{plan_text}") # Debugging

    rewritten = chain.invoke({
        "resume_text": resume_text,
        "jd_text": jd_text,
        "plan_text": plan_text,
        "role_intent": role_intent
    }).strip()

    if len(rewritten.splitlines()) < 30:
        raise RuntimeError("Rewrite failed: output too short or malformed")

    return rewritten


# =========================================================
# DOCX GENERATOR (ROBUST)
# =========================================================
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def create_docx(resume_text: str, filename: str):
    doc = Document()

    # -------------------------------------------------
    # PAGE MARGINS (Narrow-ish)
    # -------------------------------------------------
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # -------------------------------------------------
    # GLOBAL FONT (Times New Roman)
    # -------------------------------------------------
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(10.5)

    # -------------------------------------------------
    # HELPERS
    # -------------------------------------------------
    def add_bottom_border(paragraph):
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def section_header(text):
        p = doc.add_paragraph()
        run = p.add_run(text.upper())
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        
        # Add border instead of underline
        add_bottom_border(p)
        
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)

    def company_and_date(company, dates):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        table.allow_autofit = True
        
        # Ensure table doesn't have borders/grid usually, but checking style
        # Default is usually no border, checks if we need to force it off? 
        # For now, let's just set content.
        
        left = table.rows[0].cells[0]
        right = table.rows[0].cells[1]
        
        # Company Name
        p_left = left.paragraphs[0]
        r_left = p_left.add_run(company)
        r_left.bold = True
        r_left.font.name = 'Times New Roman'
        r_left.font.size = Pt(11)
        p_left.paragraph_format.space_after = Pt(0)
        
        # Date
        # If dates contains hyphen, ensure it looks nice
        p_right = right.paragraphs[0]
        r_right = p_right.add_run(dates)
        r_right.bold = True # Bold dates in target format? Image looks bold-ish.
        r_right.font.name = 'Times New Roman'
        r_right.font.size = Pt(11)
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_right.paragraph_format.space_after = Pt(0)

    def job_title(title):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.italic = True
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_after = Pt(3)

    def bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Markdown Bold Parsing: "Bold **Category:** Value"
        parts = text.split("**")
        for idx, part in enumerate(parts):
            if not part: continue
            
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            
            # Odd indices are between ** markers (e.g. [0]normal, [1]bold, [2]normal)
            if idx % 2 == 1:
                run.bold = True

    # -------------------------------------------------
    # CONTENT PROCESSING
    # -------------------------------------------------
    # -------------------------------------------------
    # CONTENT PROCESSING
    # -------------------------------------------------
    # Normalize special characters (Word bullets, en-dashes, etc.)
    resume_text = resume_text.replace("", "•").replace("–", "-").replace("—", "-")
    lines = [l.strip() for l in resume_text.split("\n") if l.strip()]
    i = 0

    # 1. HEADER SECTION (Name, Contact, Title, etc.)
    # Heuristic: The first line is Name (Large). Subsequent lines are Contact/Title (Small).
    # We continue until we hit a SECTION HEADER (All Caps) or the end.
    
    first_line = True
    while i < len(lines):
        line = lines[i]
        
        # Stop if we hit a likely Section Header (e.g. PROFESSIONAL SUMMARY)
        if line.isupper() and len(line) < 50 and not first_line:
            break
            
        # Stop if we hit a bullet (unlikely in header, but safety check)
        if line.startswith("•"):
            break

        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if first_line:
            # NAME STYLE
            run = p.runs[0]
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(20)
            p.paragraph_format.space_after = Pt(6)
            first_line = False
        else:
            # CONTACT/TITLE STYLE
            run = p.runs[0]
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            p.paragraph_format.space_after = Pt(2)
        
        i += 1
    
    # Add a little buffer after the header section before the first section content
    if i > 0 and i < len(lines):
        # We can't easily access the "last paragraph" added to add space_after, 
        # but we can check if the next line is a header and ensure it has space_before (which it does).
        pass

    # Helper for non-bullet Markdown
    # This section appears to be an accidental inclusion of prompt text within the code.
    # To maintain syntactic correctness, it's commented out.
    """
    2. AUDIT "FULL" ITEMS (Verify Evidence):
       - specific keywords must be present.
       - COMPOUND CHECK: If requirement is "Agile AND DevOps", evidence MUST show BOTH.
       - "Agile" Specifics:
         - You CANNOT infer Agile from "Scrum", "Sprints", or "Jira" alone unless the JD allows it.
         - If the JD says "Agile environment", you must find the word "Agile" or a specific methodology like "Scrum" explicitly listed.
         - IF NO EXPLICIT METHODOLOGY is named (just "Jira" or "Standups"), Downgrade to "Partial" or "Missing".
         - Note: "Inferred Agile from context" -> Status: Partial.
    """
    def markdown_paragraph(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        parts = text.split("**")
        for idx, part in enumerate(parts):
            if not part: continue
            run = p.add_run(part)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10.5)
            if idx % 2 == 1:
                run.bold = True
        p.paragraph_format.space_after = Pt(4)

    # 2. MAIN BODY LOOP
    in_experience_section = False
    
    while i < len(lines):
        line = lines[i]

        # 1. BULLETS (High Priority)
        # Handle standard bullets and Word-converted bullets
        clean_line = line.replace("•", "").replace("", "").strip()
        if line.startswith("•") or line.startswith(""):
            bullet(clean_line)
            i += 1
            continue

        # 2. SECTION HEADERS
        # More robust cleaning
        raw_line = line.strip()
        is_upper_header = raw_line.isupper() and len(raw_line) < 50 and "Location" not in raw_line
        
        # Case Insensitive Check for specific Sections
        is_known_header = any(k in raw_line.upper() for k in ["PROFESSIONAL EXPERIENCE", "WORK HISTORY", "EMPLOYMENT", "EDUCATION", "SKILLS", "CERTIFICATION", "SUMMARY", "PROJECTS"]) and len(raw_line) < 50

        if is_upper_header or is_known_header:
            is_header = True
            
            # Check if entering Experience Section (Robust Check)
            upper_line = raw_line.upper()
            if any(k in upper_line for k in ["EXPERIENCE", "WORK HISTORY", "EMPLOYMENT"]):
                in_experience_section = True
            elif any(k in upper_line for k in ["SUMMARY", "EDUCATION", "SKILL", "CERTIFICATION"]):
                in_experience_section = False
            
            # Lookahead check for Company (Date Line)
            # If next line looks like a date, this might be a Company Name, NOT a Section Header.
            # Exception: If this line explicitly says "EXPERIENCE", it IS a section header.
            if i + 1 < len(lines):
                next_l = lines[i+1].strip()
                import re
                date_start_pattern = r'^(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|^\d{4}|^Present'
                
                # If next line is a date, usually means THIS line is a Company.
                # BUT if this line is explicitly "Professional Experience", it wins.
                if re.search(date_start_pattern, next_l, re.IGNORECASE):
                    if not any(k in upper_line for k in ["EXPERIENCE", "WORK HISTORY", "EDUCATION", "SKILLS"]):
                        is_header = False
            
            if is_header:
                section_header(line.upper()) # Force Uppercase for styling
                i += 1
                continue

        # 3. COMPANY + DATE (Single Line)
        # Fix Bullet Bleed: Strip bullets before checking
        clean_text_check = line.lstrip('•-* ').strip()
        
        if any(d in clean_text_check for d in ["Present", "202", "201", "200"]) and any(sep in clean_text_check for sep in [" - ", " – ", " — "]) and len(clean_text_check) < 100:
            import re
            # Regex with Em-dash support
            match = re.search(r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|Present|\d{4})', clean_text_check)
            
            if match and match.start() > 3:
                split_idx = match.start()
                company_text = clean_text_check[:split_idx].strip() # Use clean text
                date_text = clean_text_check[split_idx:].strip()    # Use clean text
                
                company_and_date(company_text, date_text)
                i += 1
                
                if i < len(lines):
                    title_line = lines[i]
                    # Clean title too
                    clean_title = title_line.lstrip('•-* ').strip()
                    if not clean_title.isupper() and len(clean_title) < 80:
                        job_title(clean_title)
                        i += 1
                continue

        # 4. COMPANY + DATE (Two Line Pattern)
        # Only if CURRENT line is Company and NEXT is Date
        
        if i + 1 < len(lines):
            next_line_clean = lines[i+1].lstrip('•-* ').strip()
            if any(d in next_line_clean for d in ["Present", "202", "201", "200"]):
                
                # Safe to match
                has_hyphen_in_next = any(sep in next_line_clean for sep in [" - ", " – ", " — "])
                
                # Helper: Is this next line just a Date Range?
                # "Aug 2020 - Present" -> Yes. 
                # "Worked 2020 - Present" -> Maybe not. 
                # Simple check: Does it start with Month/Year/Present?
                is_date_start = re.match(r'^(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Present|\d{4})', next_line_clean, re.IGNORECASE)
                
                if is_date_start or not has_hyphen_in_next:
                    # Treat as Multi-line Company + Date
                    company_text = line.lstrip('•-* ').strip()
                    date_text = next_line_clean
                    
                    company_and_date(company_text, date_text)
                    i += 2
                    
                    if i < len(lines):
                        title_line = lines[i]
                        clean_title = title_line.lstrip('•-* ').strip()
                        if not clean_title.isupper() and len(clean_title) < 80:
                            job_title(clean_title)
                            i += 1
                    continue

            if i < len(lines):
                # Fallback: Just checking title?
                title_line = lines[i]
                clean_title = title_line.lstrip('•-* ').strip()
                if not clean_title.isupper() and not title_line.startswith("•"):
                     # Weak signal, skip
                     pass

        # DEFAULT TEXT HANDLING
        # If we are in Experience Section, assume it's a forgotten bullet
        if in_experience_section and len(line) > 10:
            bullet(line)
        else:
            # Summary or other text -> Render as Paragraph with Markdown support
            markdown_paragraph(line)
            
        i += 1
    
    doc.save(filename)
    return filename
#             p.paragraph_format.space_after = Pt(BULLET_SPACE_AFTER)

#         # NORMAL BODY TEXT
#         else:
#             p.paragraph_format.space_after = Pt(BODY_SPACE_AFTER)

#     doc.save(filename)
#     return filename



# # =========================================================
# # PDF GENERATOR (ROBUST)
# # =========================================================

# def create_pdf(text_content: str, filename: str):
#     from reportlab.lib.pagesizes import letter
#     from reportlab.pdfgen import canvas
#     from reportlab.lib.utils import simpleSplit

#     c = canvas.Canvas(filename, pagesize=letter)
#     width, height = letter
#     margin = 40
#     y = height - margin

#     c.setFont("Helvetica", 10)
#     line_height = 12

#     for line in text_content.split("\n"):
#         if y < margin:
#             c.showPage()
#             c.setFont("Helvetica", 10)
#             y = height - margin

#         indent = 20 if line.strip().startswith("•") else 0
#         wrapped = simpleSplit(
#             line.strip(),
#             "Helvetica",
#             10,
#             width - (2 * margin) - indent
#         )

#         if not wrapped:
#             y -= line_height
#             continue

#         for w in wrapped:
#             if y < margin:
#                 c.showPage()
#                 c.setFont("Helvetica", 10)
#                 y = height - margin
#             c.drawString(margin + indent, y, w)
#             y -= line_height

#     c.save()
#     return filename


# =========================================================
# COVER LETTER GENERATOR
# =========================================================

def generate_cover_letter_content(resume_text: str, jd_text: str, api_key: str) -> str:
    """Generates a tailored cover letter content."""
    
    prompt = """
    You are an Expert Career Coach and Professional Bio Writer.
    
    GOAL: Write a HIGH-IMPACT, TAILORED COVER LETTER for the Candidate based on their Resume and the Job Description.
    
    JOB DESCRIPTION:
    {jd_text}
    
    CANDIDATE RESUME:
    {resume_text}
    
    INSTRUCTIONS:
    1. Tone: Professional, Confident, Enthusiastic, but NOT robotic or cliché.
    2. Structure:
       - Header: [Check placeholders] (Name, Date, etc.)
       - Opening: Clearly state interest in the specific Role and Company (if known).
       - The "Hook": Connect the candidate's strongest achievement directly to the JD's biggest pain point.
       - The "Proof": Summarize 2-3 key skills/experiences that make them a perfect fit, using specific metrics from the resume if available.
       - Closing: Strong call to action for an interview.
    3. Formatting:
       - Use standard business letter format.
       - Do NOT use Markdown (no bolding, no headers).
       - Keep it concise (approx 300-400 words).
    
    OUTPUT:
    Return ONLY the body of the cover letter text.
    """
    
    llm = get_llm(api_key)
    
    from langchain_core.messages import HumanMessage
    try:
        res = llm.invoke([
            HumanMessage(content=prompt.format(resume_text=resume_text, jd_text=jd_text))
        ])
        content = res.content if hasattr(res, 'content') else str(res)
        return content
    except Exception as e:
        return f"Error generating cover letter: {e}"

def create_cover_letter_docx(text_content: str, filename: str):
    """Creates a simple DOCX file for the cover letter."""
    from docx import Document
    from docx.shared import Pt
    
    doc = Document()
    
    # Add content
    for line in text_content.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(12)
            
    doc.save(filename)
    return filename
